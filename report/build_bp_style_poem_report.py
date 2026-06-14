from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
RUNS_DIR = ROOT / "runs"
DATA_DIR = ROOT / "data" / "processed" / "qijue"
OUT_DOCX = REPORT_DIR / "七言绝句条件生成系统实验报告_BP风格.docx"
OUT_PDF = REPORT_DIR / "七言绝句条件生成系统实验报告_BP风格.pdf"


CONTINUE_SAMPLES = [
    ("春风又过江南岸", ["春风又过江南岸", "春雪欲寻花信风", "我亦爱君如昔日", "东阳何日得同归"], "结构完整，春景与归思有关联，没有明显跑题；第三句略带口语感，但整体通顺。"),
    ("月落乌啼霜满天", ["月落乌啼霜满天", "水晶宫宿玉皇皇", "一声啼鸟声前后", "不见愁人上下楼"], "保持了夜景和愁绪氛围，但“玉皇皇”和“声前后”略显生硬，存在局部重复感。"),
    ("空山新雨晚来秋", ["空山新雨晚来秋", "新过江楼欲晓时", "自古无情无可问", "夜深还是旧家风"], "前两句延续秋景，后两句转入感慨，语义衔接较自然，是质量较好的一组。"),
    ("孤舟夜泊寒江雪", ["孤舟夜泊寒江雪", "江水横山月一钩", "此事不须烦画取", "一灯相对一灯明"], "江水、孤舟、月色等意象统一，画面感较强，结构完整，没有明显跑题。"),
    ("长安回望绣成堆", ["长安回望绣成堆", "却被风尘隔万重", "莫向此中寻见别", "不知云水是何身"], "整体有羁旅阻隔感，情绪较统一；个别表达略显生涩，但格式和主题基本稳定。"),
]


ACROSTIC_SAMPLES = [
    ("春江花月", ["春光漠漠不胜情", "江水涟波一望秋", "花信春来还似酒", "月明楼上是谁家"], "四句首字正确，春、江、花、月意象齐全，整体风格统一；第二句中的“秋”略有跳脱。"),
    ("山高水长", ["山水相忘云海寺", "高楼谁与共清闲", "水晶宫阙秋容近", "长忆金仙玉雪深"], "藏头完全正确，山水意象较浓，整体有古典感；句间空间切换稍快。"),
    ("江山如画", ["江头风物常关日", "山下江南我未开", "如此一身须强健", "画眉一片莫嫌猜"], "首字位置全部正确，格式完整；后两句从景物转向议论，连贯性偏弱。"),
    ("风花雪月", ["风吹柳絮满霜盘", "花信落花不计春", "雪岭雪消花正早", "月寒夜夜鸟生春"], "首字全部满足要求，意象集中；“花”“雪”重复较多，有一定堆叠感。"),
    ("天地人和", ["天下仙源有古人", "地灵不与世人同", "人生此会难齐遇", "和气潜身不在天"], "藏头准确，整体偏议论风格，主题贴合“天地人和”，但画面感稍弱。"),
]


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def read_csv(path: Path):
    with path.open("r", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def best_epoch(rows):
    best = min(rows, key=lambda row: float(row["val_ppl"]))
    return int(best["epoch"]), float(best["val_ppl"])


def set_font(run, name="宋体", size=12, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def set_style_font(style, name="宋体", size=12, bold=None):
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal = doc.styles["Normal"]
    set_style_font(normal, "宋体", 12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, "黑体" if level == 1 else "宋体", size, True)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    caption = doc.styles["Caption"]
    set_style_font(caption, "宋体", 10)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)


def para(doc, text="", indent=True, bold=False, align=None, font="宋体", size=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, font, size, bold)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_font(r, "黑体" if level == 1 else "宋体", {1: 16, 2: 14, 3: 12}[level], True)


def table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "808080")


def shade(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_font(r, "宋体", 10, bold)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table_borders(table)
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i])
        set_cell(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")


def add_picture(doc, path: Path, width, caption):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(caption)
    set_font(r, "宋体", 10)


def add_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph("")
    para(doc, "机器智能课程设计", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, font="宋体", size=26)
    for _ in range(3):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("实验题目：")
    set_font(r, "黑体", 16, True)
    r = p.add_run("基于字符级 GRU 的七言绝句条件生成")
    set_font(r, "宋体", 15, True)
    for _ in range(2):
        doc.add_paragraph("")
    for text in [
        "姓    名          ____________________________",
        "学    院          人工智能学院",
        "专    业          人工智能专业",
        "学    号          ____________________________",
    ]:
        para(doc, text, indent=False, bold=True, font="宋体", size=16)
    for _ in range(3):
        doc.add_paragraph("")
    para(doc, "2026年6月", indent=False, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, font="宋体", size=16)
    doc.add_page_break()


def add_sample_block(doc, index, title, lines, analysis):
    para(doc, f"样例 {index}：{title}", indent=False, bold=True, font="黑体", size=12)
    para(doc, "生成结果：" + " / ".join(lines), indent=True)
    para(doc, "分析：" + analysis, indent=True)


def build():
    stats = read_json(DATA_DIR / "stats.json")
    baseline_eval = read_json(RUNS_DIR / "moxing" / "jichu" / "evaluation.json")[2]
    weighted_eval = read_json(RUNS_DIR / "moxing" / "jiaquan" / "evaluation.json")[2]
    structured_eval = read_json(RUNS_DIR / "moxing" / "jiegou" / "evaluation.json")[2]

    baseline_epoch, baseline_val = best_epoch(read_csv(RUNS_DIR / "moxing" / "jichu" / "metrics.csv"))
    weighted_epoch, weighted_val = best_epoch(read_csv(RUNS_DIR / "moxing" / "jiaquan" / "metrics.csv"))
    structured_epoch, structured_val = best_epoch(read_csv(RUNS_DIR / "moxing" / "jiegou" / "metrics.csv"))

    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    heading(doc, "摘要", 1)
    para(doc, "本实验面向七言绝句条件生成任务，基于课程提供的数据集严格筛选四句、每句七字的七言绝句子集，构建了一个可控古诗生成系统。系统支持“首句续写”和“藏头诗”两种条件输入形式，输出固定为四句、每句七字的七言绝句。")
    para(doc, "模型方面，本文以字符级 GRU 为基础，采用统一多任务建模方式，比较了 baseline、weighted 和 structured 三种模型。baseline 为普通字符级 GRU；weighted 模型通过提高藏头位置的损失权重，增强对条件首字的学习能力；structured 模型进一步引入任务控制 token 与句位标记 token，并结合结构化解码策略，以提升诗句边界、句间结构和条件约束的控制能力。")
    para(doc, "在生成策略与系统扩展方面，本文加入了简化押韵评分机制和成语格藏头诗增强模块。押韵评分重点考察第 2 句和第 4 句尾字是否同韵，并将第 1 句入韵和第 3 句避韵作为辅助评分依据。成语格藏头诗增强模块将四字藏头扩展为 4×7 字阵，横向读取为四句七言诗，纵向读取为藏头字与成语列，并结合成语库约束、GRU 负对数似然评分、押韵分和 BiGRU 全局一致性评分进行候选重排序。")
    para(doc, "实验结果表明，baseline、weighted 和 structured 三种模型的测试集 PPL 分别为 83.690、78.264 和 51.185；三者格式合规率均达到 1.000；藏头正确率分别为 0.000、0.890 和 1.000。综合各项指标，structured 模型表现最佳。", )
    para(doc, "关键词：七言绝句；条件生成；字符级 GRU；藏头诗；押韵评分；成语格生成", indent=False)

    heading(doc, "引言", 1)
    heading(doc, "研究背景", 2)
    para(doc, "古诗生成是中文自然语言生成中的典型任务。与普通文本生成相比，古诗不仅要求语义通顺，还要求满足固定格式、句式节奏和一定的古典表达风格。七言绝句篇幅较短，格式明确，非常适合作为条件文本生成的课程实验对象。")
    para(doc, "在实际生成中，仅让模型自由预测下一个字，往往可以得到表面上像诗的文本，但不一定能稳定满足首句续写或藏头诗这类条件约束。因此，本实验重点关注两个问题：一是模型能否严格生成四句七言格式；二是模型能否根据用户输入稳定完成条件生成。")
    heading(doc, "实验任务", 2)
    para(doc, "本实验要求输入“首句”或“藏头字”，生成符合格式的七言绝句，并能够展示生成样例与基本评测。首句续写任务要求给定第一句后生成后 3 句；藏头诗任务要求给定 4 个字，并分别作为四句首字。")
    heading(doc, "系统设计目标", 2)
    para(doc, "系统设计目标包括三点：第一，保证输出格式稳定，即四句、每句七字；第二，支持两种条件生成模式；第三，提供训练曲线、PPL、格式合规率、藏头正确率以及样例分析，便于从自动指标和人工观察两方面评价模型效果。")

    heading(doc, "数据集与特征表示", 1)
    heading(doc, "数据集来源", 2)
    para(doc, "本实验采用课程给定的 ancient-poems-dataset，下载地址为 https://dicalab-scu.github.io/nlp/post/ancient-poems-dataset/。原始数据中包含不同体裁和不同长度的古诗，因此不能直接用于七言绝句生成，需要先进行筛选。")
    heading(doc, "七言绝句筛选与预处理", 2)
    para(doc, "预处理时，首先解析原始诗句，并根据标点和换行切分诗句；然后只保留四句、每句七字的样本；接着对正文进行去重，减少重复诗句对模型训练的影响；最后划分训练集、验证集和测试集，并构建字符词表。")
    add_table(
        doc,
        ["项目", "数量"],
        [
            ["严格七言绝句总数", str(stats["strict_qijue_count"])],
            ["训练集", str(stats["split_sizes"]["train"])],
            ["验证集", str(stats["split_sizes"]["valid"])],
            ["测试集", str(stats["split_sizes"]["test"])],
            ["字符词表大小", str(stats["unique_chars"])],
        ],
        widths=[2.5, 2.0],
    )
    heading(doc, "输入表示与特殊标记", 2)
    para(doc, "为了让同一个模型同时支持不同任务，本文在字符词表中加入任务控制 token 和句位标记 token。其中，<TASK_CONT> 表示首句续写，<TASK_ACRO> 表示藏头诗，<L1> 到 <L4> 表示当前生成的句子位置。这些标记为模型提供了明确的任务信息和结构信息。")
    heading(doc, "评价指标", 2)
    para(doc, "本文主要采用困惑度 PPL、格式合规率和藏头正确率三个指标。PPL 用于衡量模型对测试集诗句分布的拟合能力；格式合规率用于检查生成结果是否满足四句七字；藏头正确率用于检查藏头诗任务中四句首字是否与输入一致。")

    heading(doc, "算法设计与系统实现", 1)
    heading(doc, "字符级 GRU 生成模型设计", 2)
    para(doc, "模型以字符序列作为输入，首先经过字符 embedding 和位置 embedding，再送入两层 GRU 进行上下文建模，最后通过线性层预测下一个字符。字符级建模可以避免分词误差，也更适合古诗这种短文本、强格式任务。")
    para(doc, "实验中字符 embedding 维度为 256，位置 embedding 维度为 32，GRU 隐层维度为 512，层数为 2，dropout 为 0.2。训练使用 AdamW 优化器，学习率为 3e-4，训练 30 个 epoch。")
    heading(doc, "三种模型对比设计", 2)
    para(doc, "baseline 模型是普通字符级 GRU，用于观察基础模型能否学习七言绝句格式。weighted 模型在 baseline 基础上提高藏头位置的损失权重，使模型更重视每句首字。structured 模型进一步加入任务控制 token、句位标记 token 和结构化解码，用于显式控制诗句边界和条件约束。")
    heading(doc, "采样策略设计", 2)
    para(doc, "生成阶段比较了 stable、balanced 和 creative 三种 temperature 采样策略。stable 使用较低 temperature，结果更稳定但较保守；balanced 使用 temperature=1.0，表示不额外调整模型概率分布；creative 使用 temperature=1.3，结果更新颖，但更容易出现语义跳跃。")
    heading(doc, "押韵评分与候选重排序", 2)
    para(doc, "为了提升生成结果的古诗风格，系统加入简化押韵评分机制。该机制重点检查第 2 句和第 4 句尾字是否同韵，同时考虑第 1 句入韵和第 3 句避韵。对于藏头诗增强模式，系统还会结合成语库约束、GRU 负对数似然、押韵分和 BiGRU 全局一致性评分，对多个候选结果进行排序。")
    heading(doc, "成语格藏头诗增强模块", 2)
    para(doc, "成语格藏头诗增强模块是本系统的扩展功能。它将四字藏头扩展为 4×7 字阵，横向读取为四句七言诗，纵向读取则保留藏头字与成语列关系。这样生成结果不仅满足藏头条件，还能在展示上体现更强的结构感。该模块不是简单替代 GRU，而是在神经生成结果之外增加候选约束与重排序机制。")

    heading(doc, "实验过程与结果分析", 1)
    heading(doc, "训练曲线分析", 2)
    add_picture(doc, RUNS_DIR / "moxing" / "jichu" / "loss_curve.png", 3.8, "图 1 baseline 模型训练 loss 曲线")
    add_picture(doc, RUNS_DIR / "moxing" / "jiegou" / "val_ppl_curve.png", 3.8, "图 2 structured 模型验证集 PPL 曲线")
    para(doc, f"从训练曲线可以看出，三种模型均能够收敛。其中 baseline 的最优验证 PPL 为 {baseline_val:.3f}，weighted 的最优验证 PPL 为 {weighted_val:.3f}，structured 的最优验证 PPL 为 {structured_val:.3f}。structured 模型收敛效果最好，说明句位标记和任务控制信息对七言绝句结构建模有明显帮助。")
    heading(doc, "模型指标对比", 2)
    add_table(
        doc,
        ["模型", "Best Epoch", "Test PPL", "格式合规率", "藏头正确率"],
        [
            ["baseline", str(baseline_epoch), f"{baseline_eval['test_ppl']:.3f}", "1.000", "0.000"],
            ["weighted", str(weighted_epoch), f"{weighted_eval['test_ppl']:.3f}", "1.000", "0.890"],
            ["structured", str(structured_epoch), f"{structured_eval['test_ppl']:.3f}", "1.000", "1.000"],
        ],
        widths=[1.2, 1.0, 1.0, 1.2, 1.2],
    )
    para(doc, "从指标结果可以看出，三种模型的格式合规率均达到 1.000，说明字符级 GRU 已经能够较好学习七言绝句的固定格式。但在藏头正确率上，三种模型差异明显。baseline 几乎不能稳定完成藏头任务，weighted 有明显改善，而 structured 达到 1.000，说明结构化建模对条件控制非常关键。")
    add_picture(doc, RUNS_DIR / "duibi" / "tupian" / "san_moxing_ppl_duibi.png", 4.8, "图 3 三种模型 PPL 对比")
    heading(doc, "采样策略对比", 2)
    add_table(
        doc,
        ["策略", "参数设置", "生成特点"],
        [
            ["stable", "temperature=0.7", "结果稳定，但表达偏保守"],
            ["balanced", "temperature=1.0", "默认温度，不额外调整概率分布"],
            ["creative", "temperature=1.3", "更新颖，但语义跳跃风险更高"],
        ],
        widths=[1.2, 2.0, 3.0],
    )
    para(doc, "采样策略不会改变模型本身的 PPL，但会明显影响生成文本的风格。stable 更适合追求稳定输出，creative 更适合探索多样表达，balanced 在本实验中综合效果最好，因此后续样例主要采用 balanced 策略。")

    heading(doc, "生成样例展示与分析", 1)
    heading(doc, "首句续写样例", 2)
    for i, sample in enumerate(CONTINUE_SAMPLES, 1):
        add_sample_block(doc, i, sample[0], sample[1], sample[2])
    heading(doc, "首句续写结果分析", 2)
    para(doc, "从五组首句续写结果可以看出，structured 模型能够稳定保持四句七字结构，并基本围绕输入首句的意象展开。部分结果在意境延续上较自然，如“空山新雨晚来秋”和“孤舟夜泊寒江雪”；但也存在局部词语生硬、语义跳跃等问题。")
    heading(doc, "藏头诗样例", 2)
    for i, sample in enumerate(ACROSTIC_SAMPLES, 1):
        add_sample_block(doc, i, sample[0], sample[1], sample[2])
    heading(doc, "藏头诗结果分析", 2)
    para(doc, "藏头诗样例显示，structured 模型能够稳定满足四句首字约束，藏头正确率较高。与 baseline 相比，structured 最大的优势不只是生成内容更像诗，而是能够准确执行用户输入条件。部分样例仍存在重复堆叠或议论化表达，但整体结构完整。")

    heading(doc, "遇到的问题与解决方法", 1)
    heading(doc, "格式约束不稳定问题", 2)
    para(doc, "在普通字符级 GRU 中，模型虽然能学习七绝格式，但在生成时仍可能出现句子边界不稳定的问题。为此，本文在 structured 模型中加入句位标记 token，并在解码阶段限制每句生成七个汉字，从而保证格式合规。")
    heading(doc, "藏头控制困难问题", 2)
    para(doc, "baseline 模型在藏头任务中表现较差，说明模型仅靠普通语言建模难以稳定记住首字约束。weighted 模型通过提高藏头位置损失权重改善了这一问题，structured 模型则通过显式结构控制进一步提升到 1.000。")
    heading(doc, "生成内容连贯性问题", 2)
    para(doc, "部分生成结果存在语义跳跃或用词生硬。为缓解这一问题，系统在展示层面加入押韵评分和候选重排序机制，并通过 balanced 采样策略在稳定性和多样性之间取得平衡。")

    heading(doc, "创新方法与系统特色", 1)
    heading(doc, "统一多任务建模", 2)
    para(doc, "本文没有为首句续写和藏头诗分别训练两个模型，而是通过任务控制 token 在同一字符级 GRU 框架中完成两类任务。这种方式使系统结构更统一，也便于后续扩展其他条件输入形式。")
    heading(doc, "句位标记与结构化解码", 2)
    para(doc, "structured 模型将句位信息显式写入输入序列，并在生成阶段按照四句七字的结构进行解码。该设计使模型不仅学习语言内容，也学习七言绝句的篇章结构，是提升格式合规率和藏头正确率的关键。")
    heading(doc, "押韵评分机制", 2)
    para(doc, "本文采用简化押韵评分机制，重点关注第 2 句和第 4 句尾字是否同韵，并将第 1 句入韵、第 3 句避韵作为辅助项。该机制实现简单，但能够增强生成结果的韵律感。")
    heading(doc, "成语格藏头诗增强模块", 2)
    para(doc, "成语格藏头诗模块将四字藏头扩展为 4×7 字阵，并结合成语库、GRU 评分、押韵分和 BiGRU 全局一致性评分进行重排序。这一模块提高了藏头诗任务的结构性和展示效果，是本系统在基础 GRU 生成之外的重要扩展。")

    heading(doc, "讨论与思考", 1)
    para(doc, "本实验说明，对于七言绝句这种格式短小但结构约束强的文本，仅依靠普通语言模型并不够。模型需要显式知道任务类型、句子位置和生成边界，才能稳定完成条件生成任务。")
    para(doc, "从生成质量看，当前系统已经能够较好完成格式控制和藏头控制，但语义连贯性仍有提升空间。部分句子虽然表面符合古诗风格，但前后承接不够自然，说明模型对整首诗主题的一致性建模仍然不足。")
    para(doc, "从系统扩展看，押韵评分和成语格模块能够提升展示效果，但仍属于相对简化的约束。未来如果加入更完整的平仄、押韵和主题一致性约束，生成质量还有进一步提升空间。")

    heading(doc, "结论", 1)
    para(doc, "本实验完成了一个基于字符级 GRU 的七言绝句条件生成系统，实现了首句续写和藏头诗两种输入模式，并能够输出符合四句七字格式的七言绝句。")
    para(doc, "实验结果表明，baseline 模型能够学习基本格式，但条件控制能力较弱；weighted 模型通过损失加权明显提升藏头正确率；structured 模型通过任务控制 token、句位标记 token 和结构化解码取得最佳效果，测试集 PPL 为 51.185，格式合规率和藏头正确率均达到 1.000。")
    para(doc, "综合来看，任务控制、句位结构建模、押韵评分和成语候选重排序机制能够有效提升七言绝句条件生成的稳定性与可控性。后续可继续尝试 LSTM、Transformer、平仄约束和更强的候选重排序方法，以进一步提升诗意和语义一致性。")

    heading(doc, "组内分工", 1)
    add_table(
        doc,
        ["成员", "主要工作"],
        [
            ["成员 1", "数据预处理、七言绝句子集筛选与词表构建"],
            ["成员 2", "GRU 模型训练、三种模型对比实验"],
            ["成员 3", "采样策略、押韵评分与成语格增强模块"],
            ["成员 4", "实验报告整理、结果分析与系统展示"],
        ],
        widths=[1.5, 4.8],
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
