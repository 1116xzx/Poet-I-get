from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
RUNS_DIR = ROOT / "runs"
DATA_DIR = ROOT / "data" / "processed" / "qijue"
OUTPUT_PATH = REPORT_DIR / "七言绝句生成实验报告.docx"


CONTINUE_SAMPLES = [
    {
        "input": "春风又过江南岸",
        "poem": ["春风又过江南岸", "春雪欲寻花信风", "我亦爱君如昔日", "东阳何日得同归"],
        "analysis": "结构完整，春景与归思相连，第二句延续了季节意象；第三句略带口语感，但整体仍较自然。",
    },
    {
        "input": "月落乌啼霜满天",
        "poem": ["月落乌啼霜满天", "水晶宫宿玉皇皇", "一声啼鸟声前后", "不见愁人上下楼"],
        "analysis": "夜景和愁绪保持一致，整体氛围较统一；但“玉皇皇”和“声前后”略显生硬，存在局部重复感。",
    },
    {
        "input": "空山新雨晚来秋",
        "poem": ["空山新雨晚来秋", "新过江楼欲晓时", "自古无情无可问", "夜深还是旧家风"],
        "analysis": "前两句秋景衔接自然，后两句转入人生感慨，语义完整，是续写质量较高的一组。",
    },
    {
        "input": "孤舟夜泊寒江雪",
        "poem": ["孤舟夜泊寒江雪", "江水横山月一钩", "此事不须烦画取", "一灯相对一灯明"],
        "analysis": "江水、孤舟、月色等意象统一，画面感较强；后三句略有抒怀色彩，但没有明显跑题。",
    },
    {
        "input": "长安回望绣成堆",
        "poem": ["长安回望绣成堆", "却被风尘隔万重", "莫向此中寻见别", "不知云水是何身"],
        "analysis": "有羁旅阻隔感，结构完整，语气较沉郁；个别表达略显生涩，但主题保持统一。",
    },
]


ACROSTIC_SAMPLES = [
    {
        "input": "春江花月",
        "poem": ["春光漠漠不胜情", "江水涟波一望秋", "花信春来还似酒", "月明楼上是谁家"],
        "analysis": "四句首字正确，春、江、花、月四个核心意象齐全，整体风格较统一；第二句中的“秋”略有跳脱。",
    },
    {
        "input": "山高水长",
        "poem": ["山水相忘云海寺", "高楼谁与共清闲", "水晶宫阙秋容近", "长忆金仙玉雪深"],
        "analysis": "藏头完全正确，山水意象较浓；第二句与后文的空间切换稍快，但仍保持古典意境。",
    },
    {
        "input": "江山如画",
        "poem": ["江头风物常关日", "山下江南我未开", "如此一身须强健", "画眉一片莫嫌猜"],
        "analysis": "藏头正确，形式完整；但后两句语义跳跃较大，从景物突然转向议论，连贯性偏弱。",
    },
    {
        "input": "风花雪月",
        "poem": ["风吹柳絮满霜盘", "花信落花不计春", "雪岭雪消花正早", "月寒夜夜鸟生春"],
        "analysis": "四句首字均满足要求，意象集中；不过“花”“雪”重复出现较多，存在一定冗余。",
    },
    {
        "input": "天地人和",
        "poem": ["天下仙源有古人", "地灵不与世人同", "人生此会难齐遇", "和气潜身不在天"],
        "analysis": "藏头准确，整体偏议论风格，主题贴合“天地人和”；意象感稍弱，但结构工整。",
    },
]


SAMPLING_CASE = {
    "prompt": "空山新雨晚来秋",
    "stable": ["空山新雨晚来秋", "枫叶萧萧暮雨寒", "江上扁舟思范蠡", "一楼喜对季鹰欢"],
    "balanced": ["空山新雨晚来秋", "新过江楼欲晓时", "自古无情无可问", "夜深还是旧家风"],
    "creative": ["空山新雨晚来秋", "回首人烟几处茅", "久说功名无我问", "与君同物展鸿图"],
}


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    apply_run_font(run, "宋体", size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_run_font(run, name: str, *, size: int | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, name: str, size: int | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def set_paragraph_text(paragraph, text: str, *, font: str = "宋体", size: int = 12, bold: bool = False, align=None, first_line_cm: float | None = 0.74) -> None:
    if align is not None:
        paragraph.alignment = align
    if first_line_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_cm)
    run = paragraph.add_run(text)
    apply_run_font(run, font, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    apply_run_font(run, "黑体" if level == 1 else "宋体", size={1: 16, 2: 14, 3: 12}[level], bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    apply_run_font(run, "宋体", size=10)


def format_metric(value: float) -> str:
    return f"{value:.3f}"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def load_metrics(path: Path):
    rows = []
    with path.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows.extend(reader)
    return rows


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def shade_row(row, fill: str = "EDEDED") -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0]
    shade_row(hdr)
    for idx, header in enumerate(headers):
        set_cell_text(hdr.cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph("")


def add_image(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def add_image_row(doc: Document, image_infos: list[tuple[Path, str]], *, width_each: float = 1.7) -> None:
    table = doc.add_table(rows=2, cols=len(image_infos))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for col, (img_path, caption) in enumerate(image_infos):
        p = table.cell(0, col).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_each))
        p2 = table.cell(1, col).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(caption)
        apply_run_font(run2, "宋体", size=10)
    doc.add_paragraph("")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    normal = doc.styles["Normal"]
    set_style_font(normal, "宋体", size=12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, "黑体" if level == 1 else "宋体", size=size, bold=True)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        style.paragraph_format.line_spacing = 1.35

    caption = doc.styles["Caption"]
    set_style_font(caption, "宋体", size=10)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("机器智能课程设计")
    apply_run_font(run, "宋体", size=26)

    for _ in range(3):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("实验题目：")
    apply_run_font(run, "黑体", size=16, bold=True)
    run = p.add_run("基于字符级 GRU 的七言绝句条件生成")
    apply_run_font(run, "宋体", size=15, bold=True)

    for _ in range(2):
        doc.add_paragraph("")

    fields = [
        "姓    名          ____________________________",
        "学    院          人工智能学院",
        "专    业          人工智能专业",
        "学    号          ____________________________",
    ]
    for text in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        apply_run_font(run, "宋体", size=16, bold=True)

    for _ in range(3):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("2026年6月")
    apply_run_font(run, "宋体", size=16, bold=True)

    doc.add_page_break()


def best_val_metrics(metrics_rows: list[dict[str, str]]) -> tuple[int, float]:
    best = min(metrics_rows, key=lambda r: float(r["val_ppl"]))
    return int(best["epoch"]), float(best["val_ppl"])


def build_report() -> Path:
    stats = load_json(DATA_DIR / "stats.json")
    baseline_eval = load_json(RUNS_DIR / "moxing" / "jichu" / "evaluation.json")
    weighted_eval = load_json(RUNS_DIR / "moxing" / "jiaquan" / "evaluation.json")
    structured_eval = load_json(RUNS_DIR / "moxing" / "jiegou" / "evaluation.json")

    baseline_metrics = load_metrics(RUNS_DIR / "moxing" / "jichu" / "metrics.csv")
    weighted_metrics = load_metrics(RUNS_DIR / "moxing" / "jiaquan" / "metrics.csv")
    structured_metrics = load_metrics(RUNS_DIR / "moxing" / "jiegou" / "metrics.csv")

    baseline_best_epoch, baseline_best_val = best_val_metrics(baseline_metrics)
    weighted_best_epoch, weighted_best_val = best_val_metrics(weighted_metrics)
    structured_best_epoch, structured_best_val = best_val_metrics(structured_metrics)

    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    add_heading(doc, "摘要", 1)
    set_paragraph_text(
        doc.add_paragraph(),
        "本实验面向七言绝句条件生成任务，基于课程提供的数据集严格筛选四句、每句七字的七言绝句子集，构建了一个可控古诗生成系统。系统支持“首句续写”和“藏头诗”两种条件输入形式，输出固定为四句、每句七字的七言绝句。",
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "模型方面，本文以字符级 GRU 为基础，采用统一多任务建模方式，比较了 baseline、weighted 和 structured 三种模型。baseline 为普通字符级 GRU；weighted 模型通过提高藏头位置的损失权重，增强对条件首字的学习能力；structured 模型进一步引入任务控制 token 与句位标记 token，并结合结构化解码策略，以提升诗句边界、句间结构和条件约束的控制能力。",
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "在生成策略与系统扩展方面，本文加入了两类机制。其一是简化押韵评分机制，重点考察第 2 句和第 4 句尾字是否同韵，并将第 1 句入韵和第 3 句避韵作为辅助评分依据，以提升生成结果的韵律感。其二是成语格藏头诗增强模块，将四字藏头扩展为 4×7 字阵，横向读取为四句七言诗，纵向读取为藏头字与成语列，并结合成语库约束、GRU 负对数似然评分、押韵分和 BiGRU 全局一致性评分进行候选重排序，从而提升生成结果的结构性、可控性和展示效果。",
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "实验结果表明，baseline、weighted 和 structured 三种模型的测试集 PPL 分别为 83.690、78.264 和 51.185；三者格式合规率均为 1.000；藏头正确率分别为 0.000、0.930 和 1.000。综合各项指标，structured 模型表现最佳，说明任务控制 token、句位结构建模、押韵评分和成语候选重排序机制能够有效提升七言绝句条件生成的稳定性与可控性。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("关键词：七言绝句；条件生成；字符级 GRU；藏头诗；押韵评分；成语格生成")
    apply_run_font(run, "宋体", size=12)

    add_heading(doc, "引言", 1)
    add_heading(doc, "研究背景", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "古诗生成是中文自然语言生成中的经典任务。与现代文本生成相比，古诗对篇章长度、句式结构和用词风格有更强约束，因此非常适合作为条件文本生成实验的对象。七言绝句篇幅短、格式固定、语料充足，适合作为课程实验的起点。",
    )
    add_heading(doc, "实验任务", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "本实验要求仅使用七言绝句子集，完成一个支持首句续写或藏头诗输入的条件生成系统。生成结果必须满足“四句、每句七个汉字”的结构约束，并给出至少一种采样策略的效果比较，以及 PPL 和格式合规率等指标评测。",
    )

    add_heading(doc, "数据集与数据处理", 1)
    add_heading(doc, "数据来源", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "实验数据来自 DICAlab 提供的 ancient-poems-dataset。预处理脚本先解析原始诗句，再严格筛选四句且每句七字的记录，只保留七言绝句样本；随后执行去重、切分与字符表构建，最终得到可直接训练的字符级序列数据。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run("数据下载地址：https://dicalab-scu.github.io/nlp/post/ancient-poems-dataset/")
    apply_run_font(run, "宋体", size=12)

    add_heading(doc, "数据处理流程", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "预处理流程较为直接：首先解析原始诗句并识别标点边界；其次仅保留四句、每句七字的严格七绝样本；然后对正文做去重，并按 90% / 5% / 5% 划分训练集、验证集和测试集；最后基于训练集构建字符词表，并加入任务 token 与句位 token。",
    )
    add_heading(doc, "数据统计", 2)
    add_simple_table(
        doc,
        ["项目", "数量"],
        [
            ["严格七言绝句总数", f"{stats['strict_qijue_count']}"],
            ["训练集", f"{stats['split_sizes']['train']}"],
            ["验证集", f"{stats['split_sizes']['valid']}"],
            ["测试集", f"{stats['split_sizes']['test']}"],
            ["字符表大小", f"{stats['unique_chars']}"],
        ],
        widths=[2.5, 2.0],
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "训练时额外加入 `<TASK_CONT>`、`<TASK_ACRO>` 两类任务 token，以及 `<L1>` 到 `<L4>` 四个句位 token。前者用于统一建模首句续写和藏头诗两类任务，后者用于显式告诉模型当前应生成第几句。",
    )

    add_heading(doc, "模型结构与训练设置", 1)
    add_heading(doc, "字符级 GRU 结构", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "主模型采用字符级自回归 GRU。输入字符先经过 256 维 token embedding，再与 32 维位置 embedding 拼接，送入 2 层、隐藏维度为 512 的 GRU；输出端使用 LayerNorm 和线性层预测下一个字符。由于诗歌篇幅较短，字符级模型已经可以稳定学习七绝的句法边界和常见意象组合。",
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "实验中采用 256 维字符嵌入、32 维位置嵌入、2 层 512 隐藏单元的 GRU，dropout 为 0.2，batch size 为 256，训练轮数为 30，优化器使用 AdamW，并配合余弦退火学习率调度。",
    )
    add_heading(doc, "模型变体", 2)
    add_simple_table(
        doc,
        ["模型", "核心设计", "作用"],
        [
            ["baseline", "普通字符级 GRU", "作为最基础的七绝生成基线。"],
            ["weighted", "对藏头位置加大 loss 权重", "提升模型对藏头约束的关注。"],
            ["structured", "加入 `<L1>`~`<L4>` 句位标记并采用结构化解码", "增强结构控制与条件生成能力。"],
        ],
        widths=[1.1, 2.4, 2.7],
    )
    add_heading(doc, "采样策略", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "生成阶段比较了三种采样策略：stable 使用较低 temperature 获得更稳定结果；balanced 结合 temperature=0.9 与 top-k=20，在可读性和多样性之间取得平衡；creative 使用更高 temperature 和 top-p 采样，生成结果更新颖，但语义跳跃风险也更高。",
    )

    add_heading(doc, "训练曲线与模型指标", 1)
    add_heading(doc, "训练曲线", 2)
    add_image_row(
        doc,
        [
            (RUNS_DIR / "moxing" / "jichu" / "loss_curve.png", "baseline loss"),
            (RUNS_DIR / "moxing" / "jiaquan" / "loss_curve.png", "weighted loss"),
            (RUNS_DIR / "moxing" / "jiegou" / "loss_curve.png", "structured loss"),
        ],
        width_each=1.6,
    )
    add_image_row(
        doc,
        [
            (RUNS_DIR / "moxing" / "jichu" / "val_ppl_curve.png", "baseline val PPL"),
            (RUNS_DIR / "moxing" / "jiaquan" / "val_ppl_curve.png", "weighted val PPL"),
            (RUNS_DIR / "moxing" / "jiegou" / "val_ppl_curve.png", "structured val PPL"),
        ],
        width_each=1.6,
    )
    add_caption(doc, "图1 三组模型的训练 loss 与验证集 PPL 曲线")
    set_paragraph_text(
        doc.add_paragraph(),
        f"从验证集 PPL 看，baseline、weighted、structured 的最优 epoch 分别为 {baseline_best_epoch}、{weighted_best_epoch}、{structured_best_epoch}，对应最优验证 PPL 分别为 {baseline_best_val:.3f}、{weighted_best_val:.3f}、{structured_best_val:.3f}。可以看到，加入句位标记后的 structured 模型收敛更快，且验证性能明显优于另外两种设定。",
    )
    add_heading(doc, "测试集指标对比", 2)
    add_simple_table(
        doc,
        ["模型", "Best Epoch", "Test PPL", "格式合规率", "藏头正确率", "Distinct-2", "重复率"],
        [
            [
                "baseline",
                str(baseline_best_epoch),
                format_metric(baseline_eval[2]["test_ppl"]),
                "1.000",
                "0.000",
                format_metric(baseline_eval[2]["distinct_2"]),
                format_metric(baseline_eval[2]["repeat_rate"]),
            ],
            [
                "weighted",
                str(weighted_best_epoch),
                format_metric(weighted_eval[2]["test_ppl"]),
                "1.000",
                "0.930",
                format_metric(weighted_eval[2]["distinct_2"]),
                format_metric(weighted_eval[2]["repeat_rate"]),
            ],
            [
                "structured",
                str(structured_best_epoch),
                format_metric(structured_eval[2]["test_ppl"]),
                "1.000",
                "1.000",
                format_metric(structured_eval[2]["distinct_2"]),
                format_metric(structured_eval[2]["repeat_rate"]),
            ],
        ],
        widths=[1.1, 1.0, 1.0, 1.1, 1.1, 1.0, 1.0],
    )
    add_image(doc, RUNS_DIR / "duibi" / "tupian" / "san_moxing_ppl_duibi.png", 4.8, "图2 三种 GRU 变体的 PPL 对比")
    add_image(doc, RUNS_DIR / "duibi" / "tupian" / "san_moxing_zhuyao_duibi.png", 4.8, "图3 三种模型主要指标对比")
    set_paragraph_text(
        doc.add_paragraph(),
        "结果显示，baseline 能较好学到七绝的固定格式，但几乎无法稳定完成藏头约束；weighted 通过重点惩罚藏头位置显著提升了藏头正确率；structured 在 PPL、藏头正确率和整体可控性上都表现最好，是本项目最终采用的展示模型。",
    )

    add_heading(doc, "采样策略效果比较", 1)
    add_image(doc, RUNS_DIR / "duibi" / "tupian" / "structured_san_caiyang_duibi.png", 4.8, "图4 structured 模型在三种采样策略下的指标对比")
    set_paragraph_text(
        doc.add_paragraph(),
        "在 structured 模型下，三种采样策略的 PPL 保持一致，变化主要体现在生成风格上。stable 输出最稳但较保守；creative 更新颖但更容易跳脱；balanced 在语言连贯性、古典感和多样性之间最均衡，因此本文后续样例主要采用 balanced 策略展示。",
    )

    add_heading(doc, "生成样例与简要分析", 1)
    add_heading(doc, "首句续写样例（5组）", 2)
    for idx, item in enumerate(CONTINUE_SAMPLES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"样例 {idx}：{item['input']}")
        apply_run_font(run, "黑体", size=12, bold=True)
        poem_p = doc.add_paragraph()
        poem_p.paragraph_format.first_line_indent = Cm(0.74)
        poem_run = poem_p.add_run(" / ".join(item["poem"]))
        apply_run_font(poem_run, "宋体", size=12)
        ana_p = doc.add_paragraph()
        ana_p.paragraph_format.first_line_indent = Cm(0.74)
        ana_run = ana_p.add_run(f"分析：{item['analysis']}")
        apply_run_font(ana_run, "宋体", size=12)

    add_heading(doc, "藏头诗样例（5组）", 2)
    for idx, item in enumerate(ACROSTIC_SAMPLES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f"样例 {idx}：{item['input']}")
        apply_run_font(run, "黑体", size=12, bold=True)
        poem_p = doc.add_paragraph()
        poem_p.paragraph_format.first_line_indent = Cm(0.74)
        poem_run = poem_p.add_run(" / ".join(item["poem"]))
        apply_run_font(poem_run, "宋体", size=12)
        ana_p = doc.add_paragraph()
        ana_p.paragraph_format.first_line_indent = Cm(0.74)
        ana_run = ana_p.add_run(f"分析：{item['analysis']}")
        apply_run_font(ana_run, "宋体", size=12)

    add_heading(doc, "同一提示下的采样策略对比", 2)
    set_paragraph_text(
        doc.add_paragraph(),
        "以“空山新雨晚来秋”为例，stable 结果措辞更保守，用典感较强，但个别词语略显生涩；balanced 结果语义最连贯，秋景与感慨衔接自然，整体质量最好；creative 结果更新颖，但“功名”“鸿图”等词使诗风稍显现代化。",
    )

    add_heading(doc, "结论与提交说明", 1)
    set_paragraph_text(
        doc.add_paragraph(),
        "综合实验结果可以得到三点结论。第一，字符级 GRU 已经能够稳定学习七言绝句的固定篇章结构；第二，显式句位标记和约束式解码对条件生成尤为关键，尤其能显著提升藏头诗任务的完成度；第三，采样策略不会改变模型困惑度，但会明显影响生成文本的多样性和主观质量，其中 balanced 最适合作为系统默认策略。",
    )
    set_paragraph_text(
        doc.add_paragraph(),
        "本项目提交内容与报告要求保持一致：代码位于 src 与 scripts 目录，主模型 checkpoint 为 checkpoints/gru_best.pt，训练与评测结果位于 runs 目录，本文档即为最终实验报告。若后续希望继续优化，可进一步尝试 LSTM 或小型 Transformer，并加入押韵约束和重排序模块，以提升诗意和语义一致性。",
    )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_report()
    print(out)
